"""Document exporter: output papers as Word (.docx) and Markdown (.md)."""

from pathlib import Path

from jinja2 import Template
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from literature_claw.writing.assistant import PaperOutline
from literature_claw.config import OUTPUT_DIR


# Markdown template
MARKDOWN_TEMPLATE = """# {{ title }}

{% for section in sections %}
## {{ section.heading }}

{{ section.content }}

{% endfor %}
"""


class Exporter:
    """Export paper outlines to Word and Markdown formats."""

    def __init__(self, output_dir: str | Path | None = None):
        self.output_dir = Path(output_dir) if output_dir else OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_markdown(self, outline: PaperOutline, filename: str = "paper.md") -> Path:
        """Export outline as a Markdown file."""
        template = Template(MARKDOWN_TEMPLATE)
        content = template.render(
            title=outline.title,
            sections=outline.sections,
        )

        output_path = self.output_dir / filename
        output_path.write_text(content, encoding="utf-8")
        return output_path

    def export_word(self, outline: PaperOutline, filename: str = "paper.docx") -> Path:
        """Export outline as a Word document with academic formatting."""
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(outline.title)
        title_run.bold = True
        title_run.font.size = Pt(16)

        doc.add_paragraph()  # Blank line after title

        # Sections
        for sec in outline.sections:
            # Section heading
            heading = doc.add_heading(sec["heading"], level=2)
            heading.style.font.name = "Times New Roman"

            # Section content
            if sec.get("content"):
                # Split content into paragraphs
                for para_text in sec["content"].split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.first_line_indent = Inches(0.5)
                        p.paragraph_format.space_after = Pt(6)

        output_path = self.output_dir / filename
        doc.save(str(output_path))
        return output_path

    def export_both(
        self, outline: PaperOutline, base_name: str = "paper"
    ) -> dict[str, Path]:
        """Export outline in both Word and Markdown formats."""
        return {
            "markdown": self.export_markdown(outline, f"{base_name}.md"),
            "word": self.export_word(outline, f"{base_name}.docx"),
        }
